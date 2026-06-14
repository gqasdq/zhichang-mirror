import io
import html
import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import streamlit as st

logger = logging.getLogger(__name__)

from agents.gold_detector.analyzer import resume_analyzer
from agents.gold_detector.reporter import report_generator
from components.reframe_compare import extract_reframe_pairs, render_reframe_compare
from components.cognitive_bias_card import render_cognitive_bias_gate, render_severe_under_hint
from components.gold_report_view import (
    render_report_body,
    render_report_header,
    render_section_divider,
    render_strengths,
)
from components.smart_navigation import (
    get_gold_detector_nav_recommendations,
    render_smart_nav,
)
from components.thinking_chain import (
    JD_MATCH_STEPS,
    RESUME_ANALYSIS_STEPS,
    run_with_thinking_chain,
)
from components.job_recommend_card import render_job_recommendations
from components.radar_compare import render_radar_compare
from components.score_ring import render_quality_ring, render_score_ring
from components.tag_badge import render_quality_tags, render_tags
from engines.jd_matcher_v2 import JDMatcherV2
from engines.job_recommender import JobRecommendation, JobRecommendResult, JobRecommender
from engines.resume_quality_scorer import ResumeQualityScorer
from ui.design_system import render_page_header
from ui.emotion_theme import apply_emotion_breath
from utils.emotion_adapter import EmotionAdapter
from core.pdf_export import export_gold_report_pdf
from core.session_manager import SessionManager
from core.analytics import track_module_enter


# ===== 鎸佷箙鍖栧瓨鍌?=====

def _probes_file() -> Path:
    return SessionManager.user_file_path("gold_probes.json")


def _load_probes() -> list[dict]:
    """浠庢枃浠跺姞杞藉巻鍙叉帰娴嬭褰?""
    probes_file = _probes_file()
    if probes_file.exists():
        try:
            data = json.loads(probes_file.read_text(encoding="utf-8"))
            if isinstance(data, list):
                return data
        except (json.JSONDecodeError, OSError):
            pass
    return []


def _save_probes(probes: list[dict]) -> None:
    """淇濆瓨鍘嗗彶鎺㈡祴璁板綍鍒版枃浠?""
    probes_file = _probes_file()
    probes_file.parent.mkdir(parents=True, exist_ok=True)
    probes_file.write_text(
        json.dumps(probes, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


# ===== 鏍峰紡 =====

def _inject_styles() -> None:
    st.markdown(
        """
<style>
/* 閲戝瓙鎺㈡祴鍣?路 椤甸潰涓撳睘鏍峰紡锛堝叏灞€鏍峰紡瑙?ui/styles.py锛?*/
[data-testid="stHorizontalBlock"] { gap: 0.75rem !important; align-items: flex-start !important; }
[data-testid="column"]:first-child { padding-left: 0 !important; }
[data-testid="stHorizontalBlock"] [data-testid="column"]:last-child {
    flex: 1 1 0 !important;
    min-width: 0 !important;
}
[data-testid="stHorizontalBlock"] [data-testid="column"]:last-child > div {
    width: 100% !important;
}
.block-container {
    max-width: 100% !important;
    width: 100% !important;
    padding-left: 12px !important;
    padding-right: 32px !important;
    padding-top: 6px !important;
}
[data-testid="stMainBlockContainer"] {
    padding-top: 6px !important;
}
.gold-report-shell,
.gold-report-shell ~ div,
div:has(> .gold-report-shell) {
    width: 100% !important;
    max-width: 100% !important;
}
h3 { margin-top: 0 !important; padding-top: 0 !important; }

/* 鎶ュ憡鏁翠綋瀹瑰櫒锛氭拺婊′富鍐呭鍒楋紝涓嶅啀閿?820px */
.gold-report-shell {
    width: 100%;
    max-width: none;
    margin: 0 0 24px;
    padding: 28px 36px 32px;
    background: rgba(255, 255, 255, 0.78);
    border: 1px solid rgba(61, 56, 51, 0.07);
    border-radius: 16px;
    box-sizing: border-box;
}

/* 鎶ュ憡澶撮儴 */
.gold-report-header {
    margin-bottom: 24px;
    padding-bottom: 20px;
    border-bottom: 1px solid rgba(61, 56, 51, 0.08);
}
.gold-report-eyebrow {
    display: inline-block;
    font-size: 12px;
    font-weight: 600;
    letter-spacing: 0.04em;
    color: #B8908A;
    margin-bottom: 6px;
}
.gold-report-title {
    font-size: 22px;
    font-weight: 650;
    color: #2C2420;
    line-height: 1.35;
    text-wrap: balance;
}
.gold-report-subtitle {
    margin-top: 6px;
    font-size: 14px;
    color: #6B5B52;
    line-height: 1.5;
}

/* 鎶ュ憡姝ｆ枃鎺掔増锛氭拺婊″鍣紝琛屽鐢?padding 鎺у埗鑰岄潪 72ch 閿佹 */
.gold-report-prose {
    width: 100%;
    max-width: none;
    color: #2C2420;
    font-size: 15px;
    line-height: 1.78;
    text-wrap: pretty;
}
.gold-prose-h2 {
    font-size: 17px;
    font-weight: 650;
    color: #2C2420;
    margin: 28px 0 14px;
    padding-bottom: 8px;
    border-bottom: 1px solid rgba(184, 144, 138, 0.22);
    text-wrap: balance;
}
.gold-prose-h3 {
    font-size: 15px;
    font-weight: 600;
    color: #3D3530;
    margin: 20px 0 10px;
}
.gold-prose-p {
    margin: 0 0 14px;
    color: #2C2420;
}
.gold-prose-p:last-child { margin-bottom: 0; }
.gold-prose-divider {
    border: none;
    border-top: 1px solid rgba(61, 56, 51, 0.08);
    margin: 24px 0;
}
.gold-prose-ol, .gold-prose-ul {
    margin: 8px 0 16px;
    padding-left: 1.35em;
}
.gold-prose-li {
    margin-bottom: 8px;
    padding-left: 4px;
    color: #2C2420;
}
.gold-report-prose strong {
    color: #2C2420;
    font-weight: 650;
}
.gold-prose-empty {
    color: #6B5B52;
    font-style: italic;
}

/* 鍖哄潡鍒嗛殧 */
.gold-section-divider {
    margin: 28px 0 20px;
    border-top: 1px solid rgba(61, 56, 51, 0.08);
    padding-top: 4px;
}
.gold-section-divider-label {
    display: inline-block;
    margin-top: -13px;
    padding: 0 10px 0 0;
    background: rgba(255, 255, 255, 0.78);
    font-size: 12px;
    font-weight: 600;
    color: #8C8279;
    letter-spacing: 0.02em;
}
.gold-section-label {
    font-size: 15px;
    font-weight: 650;
    color: #2C2420;
    margin-bottom: 14px;
}

/* 鏍稿績浼樺娍 */
.gold-strengths-block {
    margin-top: 4px;
}
.gold-strength-list {
    list-style: none;
    margin: 0;
    padding: 0;
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
    gap: 12px;
}
.gold-strength-item {
    display: flex;
    flex-direction: column;
    gap: 4px;
    padding: 14px 16px;
    background: rgba(184, 144, 138, 0.06);
    border: 1px solid rgba(184, 144, 138, 0.14);
    border-radius: 10px;
}
.gold-strength-name {
    font-size: 14px;
    font-weight: 650;
    color: #2C2420;
}
.gold-strength-desc {
    font-size: 13px;
    line-height: 1.65;
    color: #5C4F47;
}

/* 璇勫垎鎶ュ憡鍖哄潡 */
.jd-match-report {
    margin: 0;
    padding: 22px 0 0;
}
.jd-match-title {
    color: #2C2420;
    font-size: 16px;
    font-weight: 650;
    margin-bottom: 12px;
}
.jd-match-layout {
    display: flex;
    flex-wrap: wrap;
    gap: 20px;
    align-items: flex-start;
}
.jd-match-ring {
    flex: 0 0 260px;
    max-width: 280px;
}
.jd-match-dims {
    flex: 1 1 280px;
    min-width: 240px;
}
.jd-dim-row {
    margin-bottom: 14px;
}
.jd-dim-header {
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    margin-bottom: 6px;
    font-size: 13px;
    color: #2C2420;
}
.jd-dim-label { font-weight: 600; }
.jd-dim-meta { color: #8C8279; font-size: 12px; }
.jd-progress-track {
    height: 8px;
    background: rgba(61, 56, 51, 0.08);
    border-radius: 4px;
    overflow: hidden;
}
.jd-progress-fill {
    height: 100%;
    border-radius: 4px;
    transition: width 0.4s ease;
}
.jd-fill-blue { background: #4A90D9; }
.jd-fill-green { background: #5DAE8B; }
.jd-fill-orange { background: #D4956A; }
.jd-smart-suggestion {
    margin-top: 16px;
    padding: 14px 16px;
    background: rgba(184, 144, 138, 0.08);
    border-radius: 10px;
    border: 1px solid rgba(184, 144, 138, 0.18);
    color: #2C2420;
    font-size: 14px;
    line-height: 1.65;
}
.jd-smart-suggestion strong { color: #9E6B64; }

/* 瀵煎嚭涓庢搷浣滃尯 */
.gold-export-block {
    margin-top: 28px;
    padding-top: 22px;
    border-top: 1px solid rgba(61, 56, 51, 0.08);
}
.gold-export-label {
    font-size: 13px;
    font-weight: 600;
    color: #6B5B52;
    margin-bottom: 12px;
}
.gold-privacy-note {
    margin-top: 12px;
    font-size: 12px;
    color: #8C8279;
}

/* 璇︾粏鏁版嵁鎶樺彔闈㈡澘 */
.gold-report-shell [data-testid="stExpander"] {
    margin-top: 20px;
    border: 1px solid rgba(61, 56, 51, 0.08) !important;
    border-radius: 10px !important;
    background: rgba(247, 243, 239, 0.5) !important;
}
.gold-report-shell [data-testid="stExpander"] summary {
    font-size: 13px !important;
    color: #6B5B52 !important;
}

.gold-score {
    color: #B8908A;
    font-size: 34px;
    line-height: 1.2;
    font-weight: 600;
    margin: 8px 0 14px;
}

/* 鍘嗗彶闈㈡澘 - 宸﹀垪 */
.gold-history-panel {
    min-height: 0;
    background: rgba(240, 235, 227, 0.5);
    border-radius: 10px;
    padding: 8px;
}
.gold-history-divider {
    border-top: 1px solid rgba(61,56,51,0.06);
    margin: 8px 0;
}

/* 鍘嗗彶闈㈡澘鎸夐挳鏍峰紡 - 璁﹕t.button鐪嬭捣鏉ュ儚瀵硅瘽鍒楄〃椤?*/
section[data-testid="stSidebar"] button[kind="secondary"] { display: none !important; }

/* 宸﹀垪鍘嗗彶闈㈡澘閲岀殑鎸夐挳 */
.gold-history-panel button {
    text-align: left !important;
    justify-content: flex-start !important;
    font-size: 13px !important;
    padding: 6px 10px !important;
    border-radius: 8px !important;
    border: 1px solid rgba(61,56,51,0.06) !important;
    background: transparent !important;
    color: #6B5B52 !important;
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    margin-bottom: 4px !important;
    height: auto !important;
    min-height: 36px !important;
    line-height: 1.4 !important;
}
.gold-history-panel button:hover {
    background: rgba(184, 144, 138, 0.08) !important;
    border-color: rgba(184, 144, 138, 0.2) !important;
}

/* 杩介棶鍒嗘瀽甯?*/
.gold-followup-wrap {
    margin-top: 28px;
    padding: 22px 24px;
    background: rgba(255, 255, 255, 0.55);
    border: 1px solid rgba(61, 56, 51, 0.07);
    border-radius: 14px;
}
.gold-followup-title {
    color: #2C2420;
    font-size: 15px;
    font-weight: 650;
    margin-bottom: 14px;
}
[class^="st-key-gold_fq_"] button {
    background: rgba(184, 144, 138, 0.06) !important;
    color: #6B5B52 !important;
    border: 1px solid rgba(184, 144, 138, 0.35) !important;
    border-radius: 8px !important;
    font-size: 13px !important;
    font-weight: 500 !important;
    min-height: 38px !important;
}
[class^="st-key-gold_fq_"] button:hover {
    background: rgba(184, 144, 138, 0.14) !important;
    border-color: #B8908A !important;
}
.st-key-gold_followup_send button[kind="primary"] {
    background-color: #B8908A !important;
    border-color: #B8908A !important;
    color: #FFF !important;
}
.st-key-gold_followup_send button[kind="primary"]:hover {
    background-color: #A07A74 !important;
    border-color: #A07A74 !important;
}
</style>
""",
        unsafe_allow_html=True,
    )


# ===== 鐘舵€佸垵濮嬪寲 =====

def _init_state() -> None:
    if "gold_resume_text" not in st.session_state:
        st.session_state.gold_resume_text = ""
    if "gold_jd_text" not in st.session_state:
        st.session_state.gold_jd_text = ""
    if "gold_conversations" not in st.session_state:
        # 浠庢寔涔呭寲鏂囦欢鍔犺浇鍘嗗彶璁板綍
        st.session_state.gold_conversations = _load_probes()
    if "gold_current_id" not in st.session_state:
        st.session_state.gold_current_id = None
    if "gold_current_conv_id" not in st.session_state:
        st.session_state.gold_current_conv_id = None
    if "gold_current_result" not in st.session_state:
        st.session_state.gold_current_result = None
    if "gold_show_input" not in st.session_state:
        st.session_state.gold_show_input = True
    if "gold_pending_probe" not in st.session_state:
        st.session_state.gold_pending_probe = None
    if "gold_probe_running" not in st.session_state:
        st.session_state.gold_probe_running = False
    if "gold_flash_message" not in st.session_state:
        st.session_state.gold_flash_message = None
    if "gold_flash_type" not in st.session_state:
        st.session_state.gold_flash_type = "info"
    if "gold_upload_name" not in st.session_state:
        st.session_state.gold_upload_name = None
    if "gold_resume_input" not in st.session_state:
        st.session_state.gold_resume_input = st.session_state.gold_resume_text
    if "gold_jd_input" not in st.session_state:
        st.session_state.gold_jd_input = st.session_state.gold_jd_text
    if "gold_jd_list" not in st.session_state:
        existing_jd = st.session_state.get("gold_jd_text") or st.session_state.get("gold_jd_input", "")
        st.session_state.gold_jd_list = [{"name": "宀椾綅1", "content": existing_jd}]
    if "gold_followup_history" not in st.session_state:
        st.session_state.gold_followup_history = {}
    if "gold_followup_pending" not in st.session_state:
        st.session_state.gold_followup_pending = None


# ===== 鎶ュ憡瑙ｆ瀽 =====

def _parse_report_text(result: dict) -> str:
    raw_content = result.get("report", {}).get("raw_content", "")
    return parse_report(raw_content)


def parse_report(raw_content: str) -> str:
    """浠巖eporter鐨剅aw_content涓彁鍙栬嚜鐒惰瑷€鎶ュ憡"""
    text = (raw_content or "").strip()
    if "```json" in text:
        text = text.split("```json")[1].split("```")[0].strip()
    elif "```" in text:
        text = text.split("```")[1].split("```")[0].strip()

    try:
        parsed = json.loads(text)
        if "natural_language_report" in parsed:
            return parsed["natural_language_report"]

        parts = []
        if "opening" in parsed:
            parts.append(parsed["opening"])
        if "your_golds" in parsed:
            for gold in parsed["your_golds"]:
                title = gold.get("title", "")
                desc = gold.get("description", "")
                why = gold.get("why_valuable", "")
                parts.append(f" **{title}** \n{desc}\n{why}")
        if "hidden_sparkles" in parsed:
            for sp in parsed["hidden_sparkles"]:
                title = sp.get("title", "")
                value = sp.get("real_value", "")
                parts.append(f" **{title}** \n{value}")
        if "gap_reframes" in parsed:
            for gap in parsed["gap_reframes"]:
                parts.append(f"宸窛锛歿gap.get('gap', '')}\n缈绘锛歿gap.get('reframe', '')}")
        if "next_actions" in parsed:
            parts.append("**涓嬩竴姝ヨ鍔?*")
            for action in parsed["next_actions"]:
                parts.append(f"- {action.get('action', '')}锛歿action.get('how', '')}")
        if "closing" in parsed:
            parts.append(parsed["closing"])
        if parts:
            return "\n\n".join(parts)
    except json.JSONDecodeError:
        pass

    nl_match = re.search(r'"natural_language_report"\s*:\s*"((?:[^"\\]|\\.)*)"', raw_content)
    if nl_match:
        return nl_match.group(1).replace("\\n", "\n").replace('\\"', '"')

    # 鍏滃簳锛氬鏋滆繑鍥炲唴瀹硅繕鏄疛SON寮€澶村氨鍐嶅墺涓€灞?    if text.startswith("{"):
        try:
            parsed = json.loads(text)
            if "natural_language_report" in parsed:
                return parsed["natural_language_report"]
        except json.JSONDecodeError:
            pass
    return raw_content


def _parse_match_score(result: dict):
    match_data = _parse_match_data(result)
    if match_data and match_data.get("overall_score") is not None:
        return int(round(float(match_data["overall_score"])))
    match_raw = result.get("match", {}).get("raw_content", "") if result.get("match") else ""
    score_match = re.search(r'"match_score"\s*:\s*(\d+)', match_raw)
    if score_match:
        return int(score_match.group(1))
    score_match = re.search(r'"overall_score"\s*:\s*([\d.]+)', match_raw)
    return int(round(float(score_match.group(1)))) if score_match else None


def _get_match_results(result: dict) -> list[dict]:
    """浠庣粨鏋滀腑瑙ｆ瀽澶?JD 鍖归厤鍒楄〃锛屽吋瀹规棫鍗?JD 鏍煎紡銆?""
    match_results = result.get("match_results")
    if not match_results and result.get("match"):
        return [{"name": "鐩爣宀椾綅", "result": result["match"]}]
    return match_results or []


def _parse_match_data(result: dict) -> Optional[dict]:
    """浠庣粨鏋滀腑瑙ｆ瀽 JD 涓夌淮鍖归厤鏁版嵁銆?""
    match = result.get("match")
    if not match:
        return None

    # 鏃х増璇瓨鍏?match 鐨勬棤 JD 鏁版嵁锛屽拷鐣?    if match.get("has_jd") is False:
        return None

    if match.get("keyword_score") is not None or match.get("overall_score") is not None:
        return match

    raw = match.get("raw_content", "")
    if not raw:
        return None

    cleaned = raw.strip()
    if "```json" in cleaned:
        cleaned = cleaned.split("```json", 1)[1].split("```", 1)[0].strip()
    elif "```" in cleaned:
        cleaned = cleaned.split("```", 1)[1].split("```", 1)[0].strip()

    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict) and (
            "keyword_score" in parsed or "overall_score" in parsed or "match_score" in parsed
        ):
            return parsed
    except json.JSONDecodeError:
        pass

    return None


def _parse_quality_data(result: dict) -> Optional[dict]:
    """浠庣粨鏋滀腑瑙ｆ瀽绠€鍘嗚川閲忚瘎浼版暟鎹€?""
    quality = result.get("quality")
    if not quality:
        return None

    if quality.get("star_score") is not None or quality.get("overall_score") is not None:
        return quality

    raw = quality.get("raw_content", "")
    if not raw:
        return None

    cleaned = raw.strip()
    if "```json" in cleaned:
        cleaned = cleaned.split("```json", 1)[1].split("```", 1)[0].strip()
    elif "```" in cleaned:
        cleaned = cleaned.split("```", 1)[1].split("```", 1)[0].strip()

    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict) and (
            "star_score" in parsed or "expression_score" in parsed or "overall_score" in parsed
        ):
            return parsed
    except json.JSONDecodeError:
        pass

    return None


def _render_dimension_bar(label: str, score: int, meta: str, color_class: str) -> str:
    safe_label = html.escape(label, quote=True)
    safe_meta = html.escape(meta, quote=True)
    pct = max(0, min(100, int(score)))
    return f"""
<div class="jd-dim-row">
  <div class="jd-dim-header">
    <span class="jd-dim-label">{safe_label} {pct}%</span>
    <span class="jd-dim-meta">{safe_meta}</span>
  </div>
  <div class="jd-progress-track">
    <div class="jd-progress-fill {color_class}" style="width:{pct}%"></div>
  </div>
</div>
"""


def _ensure_quality_data(
    result: dict,
    conv_id: Optional[str] = None,
    jd: str = "",
) -> Optional[dict]:
    """涓烘棫璁板綍琛ュ叏绠€鍘嗚川閲忔暟鎹紙鏃?JD 涓旀棤 quality 瀛楁鏃讹級銆?""
    quality_data = _parse_quality_data(result)
    if quality_data:
        return quality_data

    if _parse_match_data(result):
        return None

    if (jd or "").strip():
        return None

    analysis_raw = (result.get("analysis") or {}).get("raw_content", "")
    if not analysis_raw:
        return None

    cache_key = f"gold_backfill_quality_{conv_id or 'current'}"
    if cache_key in st.session_state:
        return st.session_state[cache_key]

    try:
        def _eval_quality():
            return ResumeQualityScorer().evaluate(analysis_raw)

        quality_result = run_with_thinking_chain(
            RESUME_ANALYSIS_STEPS,
            _eval_quality,
            model_name="DeepSeek V3 路 鍒嗘瀽鎺ㄧ悊",
        )
        quality_data = quality_result.model_dump()
        st.session_state[cache_key] = quality_data

        if conv_id:
            for conv in st.session_state.get("gold_conversations", []):
                if conv.get("id") == conv_id:
                    conv.setdefault("result", {})["quality"] = quality_data
                    if conv["result"].get("match", {}).get("has_jd") is False:
                        conv["result"]["match"] = None
                    _save_probes(st.session_state.gold_conversations)
                    break

        return quality_data
    except Exception as e:
        logger.warning("[gold_detector] backfill quality failed: %s", e)
        return None


def _render_match_report(
    match_data: dict,
    report_key: str = "default",
    result: Optional[dict] = None,
) -> None:
    """娓叉煋宀椾綅鍖归厤鎶ュ憡锛堟湁 JD 妯″紡锛夈€?""
    adapter = EmotionAdapter.from_session()
    layout = adapter.get_layout_mode()
    theme = adapter.get_theme()
    progress_style = adapter.get_progress_style()

    keyword_score = int(match_data.get("keyword_score", 0))
    star_score = int(match_data.get("star_score", 0))
    quant_score = int(match_data.get("quant_score", 0))
    overall_score = float(match_data.get("overall_score", 0))

    matched = match_data.get("keyword_matched") or []
    missing = match_data.get("keyword_missing") or []
    star_details = match_data.get("star_details") or []
    quant_details = match_data.get("quant_details") or []
    smart_suggestion = (match_data.get("smart_suggestion") or "").strip()

    star_pending = [
        item for item in star_details
        if isinstance(item, dict) and str(item.get("status", "")).lower() not in {"complete", "ok", "done"}
    ]
    quant_ok = sum(
        1 for item in quant_details
        if isinstance(item, dict) and str(item.get("status", "")).lower() == "quantified"
    )
    quant_pending = sum(
        1 for item in quant_details
        if isinstance(item, dict) and str(item.get("status", "")).lower() != "quantified"
    )

    dimensions = [
        ("鍏抽敭璇嶅尮閰?, keyword_score, f"宸插尮閰?{len(matched)} / 寰呰ˉ鍏?{len(missing)}", "jd-fill-blue"),
        ("STAR 缁撴瀯", star_score, f"{len(star_pending)} 娈靛緟鏀瑰啓" if star_pending else "缁撴瀯瀹屾暣", "jd-fill-green"),
        ("閲忓寲琛ㄨ揪", quant_score, f"{quant_ok} 澶勫凡閲忓寲 / {quant_pending} 澶勫緟閲忓寲", "jd-fill-orange"),
    ]

    if not render_cognitive_bias_gate(int(round(overall_score)), report_key=report_key):
        return

    from engines.cognitive_bias_detector import detect_cognitive_bias, should_show_bias_detection

    emotion_raw = (
        st.session_state.get("workshop_emotion_state")
        or st.session_state.get("emotion_state")
        or "骞崇ǔ"
    )
    if should_show_bias_detection(str(emotion_raw)):
        self_key = f"gold_self_match_{report_key}"
        if st.session_state.get(f"gold_bias_revealed_{report_key}"):
            self_score = int(st.session_state.get(self_key, 30))
            bias = detect_cognitive_bias(self_score, int(round(overall_score)), str(emotion_raw))
            if bias.severity == "涓ラ噸浣庝及":
                render_severe_under_hint()
        st.markdown(
            '<div style="border-top:1px solid rgba(61,56,51,0.08);margin:20px 0;"></div>',
            unsafe_allow_html=True,
        )

    if layout == "guided":
        adapter.render_guided_steps(
            ["鈶?鍏堢湅鎬诲垎", "鈶?鍐嶇湅姣忎釜缁村害鐨勮鎯?, "鈶?鏈€鍚庣湅 AI 寤鸿"],
            title="鍒€ワ紝璺熺潃鐪?,
        )
    elif layout == "praise_first":
        dimensions.sort(key=lambda x: x[1], reverse=True)

    st.markdown(f'<div class="{adapter.get_shell_class("jd-match-report")}">', unsafe_allow_html=True)
    title_prefix = theme.get("emoji_prefix", "馃幆")
    if layout == "praise_first":
        st.markdown(f'<div class="jd-match-title">{title_prefix} 鍏堢湅鐪嬩綘鍋氬緱濂界殑</div>', unsafe_allow_html=True)
    elif layout == "single_column":
        st.markdown(f'<div class="jd-match-title">{title_prefix} 浣犵殑绠€鍘嗗尮閰嶅害</div>', unsafe_allow_html=True)
        st.caption("馃挕 缁煎悎鍒嗗凡涓轰綘淇濈暀锛岀粏椤圭粏鑺傛敹鍦ㄤ笅鏂癸紝闇€瑕佹椂鍐嶅睍寮€銆?)
    else:
        st.markdown('<div class="jd-match-title">馃幆 宀椾綅鍖归厤鎶ュ憡</div>', unsafe_allow_html=True)

    ring_col, dims_col = st.columns([1, 2])
    with ring_col:
        render_score_ring(overall_score, keyword_score, star_score, quant_score)

    show_dims_inline = layout != "single_column"
    with dims_col:
        if show_dims_inline:
            bars_html = ""
            for idx, (label, score, meta, color_class) in enumerate(dimensions):
                if layout == "praise_first" and idx == 0:
                    bars_html += f'<div class="highlight-done" style="padding:4px 0; margin-bottom:4px;">'
                else:
                    bars_html += ""
                dim_meta = meta if progress_style != "minimal" else ""
                bars_html += _render_dimension_bar(label, score, dim_meta, color_class)
                if layout == "praise_first" and idx == 0:
                    bars_html += "</div>"
            st.markdown(f'<div class="jd-match-dims">{bars_html}</div>', unsafe_allow_html=True)
        else:
            st.caption("缁煎悎鍖归厤搴﹀凡涓轰綘绠€鍖栧睍绀猴紝缁嗚妭鍙互绋嶅悗鍐嶇湅銆?)

    if layout == "single_column":
        with st.expander("鏌ョ湅鍚勭淮搴﹁鎯?):
            bars_html = "".join(
                _render_dimension_bar(label, score, meta, color_class)
                for label, score, meta, color_class in dimensions
            )
            st.markdown(f'<div class="jd-match-dims">{bars_html}</div>', unsafe_allow_html=True)
            render_tags(matched, missing, star_pending)
    else:
        render_tags(matched, missing, star_pending)

    if smart_suggestion:
        safe_suggestion = html.escape(smart_suggestion, quote=True)
        st.markdown(
            f'<div class="jd-smart-suggestion"><strong>馃挕 鏅鸿兘寤鸿锛?/strong>{safe_suggestion}</div>',
            unsafe_allow_html=True,
        )

    if st.button(
        "杩涘叆閲戝瓙宸ュ潑锛屽紑濮嬩紭鍖?鈫?,
        type="primary",
        use_container_width=True,
        key=f"gold_workshop_{report_key}",
    ):
        _navigate_to_workshop(_get_workshop_result(match_data=match_data, quality_data=None) if result is None else result)

    st.markdown("</div>", unsafe_allow_html=True)


def _render_quality_report(
    quality_data: dict,
    report_key: str = "default",
    result: Optional[dict] = None,
) -> None:
    """娓叉煋绠€鍘嗚川閲忔姤鍛婏紙鏃?JD 妯″紡锛夈€?""
    adapter = EmotionAdapter.from_session()
    layout = adapter.get_layout_mode()
    theme = adapter.get_theme()
    progress_style = adapter.get_progress_style()

    star_score = int(quality_data.get("star_score", 0))
    quant_score = int(quality_data.get("quant_score", 0))
    expression_score = int(quality_data.get("expression_score", 0))
    overall_score = float(quality_data.get("overall_score", 0))

    star_details = quality_data.get("star_details") or []
    quant_details = quality_data.get("quant_details") or []
    expression_details = quality_data.get("expression_details") or []
    quality_suggestion = (quality_data.get("quality_suggestion") or "").strip()

    star_pending = [
        item for item in star_details
        if isinstance(item, dict) and str(item.get("status", "")).lower() not in {"complete", "ok", "done"}
    ]
    quant_pending = [
        item for item in quant_details
        if isinstance(item, dict) and str(item.get("status", "")).lower() != "quantified"
    ]
    expression_pending = [
        item for item in expression_details
        if isinstance(item, dict) and str(item.get("status", "")).lower() not in {"complete", "ok", "done"}
    ]

    colloquial_count = sum(
        1 for item in expression_pending
        if str(item.get("status", "")).lower() == "colloquial"
    )
    hollow_count = sum(
        1 for item in expression_pending
        if str(item.get("status", "")).lower() in {"hollow_claim", "hollow", "empty_claim"}
    )
    expression_meta_parts = []
    if colloquial_count:
        expression_meta_parts.append(f"{colloquial_count} 澶勫彛璇寲")
    if hollow_count:
        expression_meta_parts.append(f"{hollow_count} 澶勭┖璇濆璇?)
    if not expression_meta_parts:
        expression_meta_parts.append("琛ㄨ揪瑙勮寖" if not expression_pending else f"{len(expression_pending)} 澶勫緟浼樺寲")
    expression_meta = " / ".join(expression_meta_parts)

    if star_pending:
        total = max(len(star_details), len(star_pending))
        star_meta = f"{total} 娈电粡鍘嗕腑 {len(star_pending)} 娈靛緟鏀瑰啓"
    else:
        star_meta = "缁撴瀯瀹屾暣"

    dimensions = [
        ("STAR 缁撴瀯", star_score, star_meta, "jd-fill-green"),
        (
            "閲忓寲琛ㄨ揪",
            quant_score,
            f"{len(quant_details) - len(quant_pending)} 澶勫凡閲忓寲 / {len(quant_pending)} 澶勫緟閲忓寲",
            "jd-fill-orange",
        ),
        ("琛ㄨ揪瑙勮寖", expression_score, expression_meta, "jd-fill-blue"),
    ]

    if layout == "guided":
        adapter.render_guided_steps(
            ["鈶?鍏堢湅鎬诲垎", "鈶?鍐嶇湅姣忎釜缁村害鐨勮鎯?, "鈶?鏈€鍚庣湅 AI 寤鸿"],
            title="鍒€ワ紝璺熺潃鐪?,
        )
    elif layout == "praise_first":
        dimensions.sort(key=lambda x: x[1], reverse=True)

    st.markdown(f'<div class="{adapter.get_shell_class("jd-match-report")}">', unsafe_allow_html=True)
    title_prefix = theme.get("emoji_prefix", "馃搵")
    if layout == "praise_first":
        st.markdown(f'<div class="jd-match-title">{title_prefix} 鍏堢湅鐪嬩綘鍋氬緱濂界殑</div>', unsafe_allow_html=True)
    elif layout == "single_column":
        st.markdown(f'<div class="jd-match-title">{title_prefix} 浣犵殑绠€鍘嗚川閲?/div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="jd-match-title">馃搵 绠€鍘嗚川閲忔姤鍛?/div>', unsafe_allow_html=True)
    st.caption("鏈～鍐欏矖浣?JD锛屼互涓嬩负绠€鍘嗚嚜韬川閲忚瘎浼般€傜矘璐寸洰鏍囧矖浣?JD 鍚庡彲瑙ｉ攣鍏抽敭璇嶅尮閰嶅垎鏋愩€?)

    ring_col, dims_col = st.columns([1, 2])
    with ring_col:
        render_quality_ring(overall_score, star_score, quant_score, expression_score)

    show_dims_inline = layout != "single_column"
    with dims_col:
        if show_dims_inline:
            bars_html = ""
            for idx, (label, score, meta, color_class) in enumerate(dimensions):
                if layout == "praise_first" and idx == 0:
                    bars_html += '<div class="highlight-done" style="padding:4px 0; margin-bottom:4px;">'
                dim_meta = meta if progress_style != "minimal" else ""
                bars_html += _render_dimension_bar(label, score, dim_meta, color_class)
                if layout == "praise_first" and idx == 0:
                    bars_html += "</div>"
            st.markdown(f'<div class="jd-match-dims">{bars_html}</div>', unsafe_allow_html=True)
        else:
            st.caption("缁煎悎璐ㄩ噺宸蹭负浣犵畝鍖栧睍绀猴紝缁嗚妭鍙互绋嶅悗鍐嶇湅銆?)

    if layout == "single_column":
        with st.expander("鏌ョ湅鍚勭淮搴﹁鎯?):
            bars_html = "".join(
                _render_dimension_bar(label, score, meta, color_class)
                for label, score, meta, color_class in dimensions
            )
            st.markdown(f'<div class="jd-match-dims">{bars_html}</div>', unsafe_allow_html=True)
            render_quality_tags(star_pending, quant_pending, expression_pending)
    else:
        render_quality_tags(star_pending, quant_pending, expression_pending)

    if quality_suggestion:
        safe_suggestion = html.escape(quality_suggestion, quote=True)
        st.markdown(
            f'<div class="jd-smart-suggestion"><strong>馃挕 璐ㄩ噺寤鸿锛?/strong>{safe_suggestion}</div>',
            unsafe_allow_html=True,
        )

    star_evidence = quality_data.get("star_evidence") or []
    quant_evidence = quality_data.get("quant_evidence") or []
    expression_evidence = quality_data.get("expression_evidence") or []
    has_xai = any([star_evidence, quant_evidence, expression_evidence])

    if has_xai or star_pending or quant_pending or expression_pending:
        with st.expander("馃攳 鏌ョ湅璇勫垎渚濇嵁锛堝彲瑙ｉ噴 AI / XAI锛?, expanded=has_xai):
            from components.xai_evidence import render_xai_evidence_section

            render_xai_evidence_section(
                "STAR 缁撴瀯",
                star_evidence or [
                    {
                        "original_text": item.get("original_text", item.get("content", "")),
                        "issue": item.get("status", ""),
                        "suggestion": item.get("suggestion", ""),
                    }
                    for item in star_pending
                    if isinstance(item, dict)
                ],
            )
            render_xai_evidence_section(
                "閲忓寲琛ㄨ揪",
                quant_evidence or [
                    {
                        "original_text": item.get("original_text", item.get("content", "")),
                        "issue": item.get("status", ""),
                        "suggestion": item.get("suggestion", ""),
                    }
                    for item in quant_pending
                    if isinstance(item, dict)
                ],
            )
            render_xai_evidence_section(
                "琛ㄨ揪瑙勮寖",
                expression_evidence or [
                    {
                        "original_text": item.get("original_text", item.get("content", "")),
                        "issue": item.get("status", ""),
                        "suggestion": item.get("suggestion", ""),
                    }
                    for item in expression_pending
                    if isinstance(item, dict)
                ],
            )

    if st.button("绮樿创 JD锛岃В閿佸叧閿瘝鍖归厤 鈫?, key=f"gold_unlock_jd_{report_key}"):
        st.session_state.gold_show_input = True
        st.rerun()

    if st.button(
        "杩涘叆閲戝瓙宸ュ潑锛屽紑濮嬩紭鍖?鈫?,
        type="primary",
        use_container_width=True,
        key=f"gold_workshop_quality_{report_key}",
    ):
        _navigate_to_workshop(_get_workshop_result(match_data=None, quality_data=quality_data) if result is None else result)

    st.markdown("</div>", unsafe_allow_html=True)


def _extract_strengths(result: dict, report_text: str) -> list[str]:
    """浠巃nalyzer鐨勫垎鏋愮粨鏋滀腑绮惧噯鎻愬彇鏍稿績浼樺娍锛屼笉渚濊禆reporter鐨勮嚜鐒惰瑷€鎶ュ憡"""
    strengths: list[str] = []

    # 绗竴姝ワ細浠巃nalysis鐨剅aw_content涓彁鍙朿ore_advantages
    analysis_raw = result.get("analysis", {}).get("raw_content", "")
    analysis_cleaned = analysis_raw.strip()
    if "```json" in analysis_cleaned:
        analysis_cleaned = analysis_cleaned.split("```json")[1].split("```")[0].strip()
    elif "```" in analysis_cleaned:
        analysis_cleaned = analysis_cleaned.split("```")[1].split("```")[0].strip()

    try:
        parsed_analysis = json.loads(analysis_cleaned)
        for gold in parsed_analysis.get("core_advantages", []):
            if not isinstance(gold, dict):
                continue
            name = (gold.get("name") or gold.get("title") or gold.get("advantage") or "").strip()
            diff = (gold.get("differentiation") or gold.get("description") or gold.get("detail") or "").strip()
            market = (gold.get("market_value") or "").strip()
            # 缁勫悎锛氬悕绉?+ 宸紓鍖栨弿杩?            if name and diff:
                strengths.append(f"{name}锛歿diff}")
            elif name:
                strengths.append(name)
    except json.JSONDecodeError:
        pass

    # 绗簩姝ワ細濡傛灉analysis閲屾病鏈夛紝灏濊瘯浠巖eport鐨剅aw_content閲屾壘
    if not strengths:
        raw = result.get("report", {}).get("raw_content", "")
        cleaned = raw.strip()
        if "```json" in cleaned:
            cleaned = cleaned.split("```json")[1].split("```")[0].strip()
        elif "```" in cleaned:
            cleaned = cleaned.split("```")[1].split("```")[0].strip()
        try:
            parsed = json.loads(cleaned)
            # reporter鍙兘鎶奱nalysis缁撴灉甯﹁繘鏉ヤ簡
            for gold in parsed.get("core_advantages", []):
                if not isinstance(gold, dict):
                    continue
                name = (gold.get("name") or gold.get("title") or "").strip()
                diff = (gold.get("differentiation") or gold.get("description") or "").strip()
                if name and diff:
                    strengths.append(f"{name}锛歿diff}")
                elif name:
                    strengths.append(name)
        except json.JSONDecodeError:
            pass

    # 绗笁姝ワ細姝ｅ垯鍏滃簳锛屼粠analysis_raw涓尮閰峮ame瀛楁
    if not strengths:
        arr_match = re.search(r'"core_advantages"\s*:\s*\[(.*?)\]', analysis_raw, re.DOTALL)
        if arr_match:
            name_matches = re.findall(r'"name"\s*:\s*"((?:[^"\\]|\\.)*?)"', arr_match.group(1))
            for n in name_matches:
                n = n.replace("\\n", "\n").replace('\\"', '"').strip()
                if n and n not in strengths:
                    strengths.append(n)

    # 娓呯悊鍜屽幓閲?    cleaned: list[str] = []
    for item in strengths:
        text = item.strip()
        if not text or text in {"-", "鈥?, "*", "锛?, "**"}:
            continue
        # 杩囨护鎺夋柟娉曡鐗瑰緛鐨勫唴瀹?        method_keywords = ["璁拌处寮?, "STAR娉曞垯", "鍔ㄨ瘝+浠诲姟+缁撴灉", "閿欒绀鸿寖", "姝ｇ‘绀鸿寖", "绗竴姝?, "绗簩姝?, "绗笁姝?, "鎬庝箞鏀?, "鏀规硶"]
        if any(kw in text for kw in method_keywords):
            continue
        if text not in cleaned:
            cleaned.append(text)

    return cleaned[:5]


_FOLLOWUP_SYSTEM_PROMPT = (
    '浣犳槸涓€浣嶇畝鍘嗗垎鏋愬笀"閲戝瓙"锛岀敤鎴峰浣犵殑鎶ュ憡鏈夌枒闂紝璇峰熀浜庣畝鍘嗗師鏂囧拰鎶ュ憡鍐呭鍥炵瓟銆?
    '濡傛灉娑夊強"鎬庝箞琛ㄨ揪"锛岀粰鍑哄師鏂噕s鏀瑰啓瀵规瘮锛圫TAR娉曞垯+閲忓寲鎴愭灉锛夈€?
    "濡傛灉娑夊強鐭澘锛岀粰鍑哄叿浣撹ˉ鏁戣矾绾裤€?
)

_FOLLOWUP_QUICK_QUESTIONS = [
    "浣犺鐨勪紭鍔垮叿浣撳摢閲屼綋鐜帮紵",
    "杩欎簺鑳藉姏姹傝亴鏃舵€庝箞琛ㄨ揪锛?,
    "鎴戞渶澶х殑鐭澘鏄粈涔堬紵鎬庝箞琛ワ紵",
]


def _get_current_conv() -> Optional[dict]:
    conv_id = st.session_state.get("gold_current_conv_id")
    if not conv_id:
        return None
    return next((c for c in st.session_state.gold_conversations if c["id"] == conv_id), None)


def _get_followup_conv_key() -> str:
    return st.session_state.get("gold_current_conv_id") or "current"


def _get_followup_history() -> list[dict]:
    key = _get_followup_conv_key()
    history_map = st.session_state.gold_followup_history
    if key not in history_map:
        history_map[key] = []
    return history_map[key]


def _get_resume_for_followup() -> str:
    resume = (st.session_state.get("gold_resume_text") or "").strip()
    if resume:
        return resume
    conv = _get_current_conv()
    if conv:
        return (conv.get("resume") or "").strip()
    return ""


def _emotion_state_for_workshop() -> str:
    from utils.emotion_adapter import sync_emotion_to_session

    return sync_emotion_to_session()


def _get_workshop_result(
    match_data: Optional[dict] = None,
    quality_data: Optional[dict] = None,
) -> dict:
    """鑾峰彇褰撳墠鎶ュ憡瀹屾暣 result锛屼緵璺宠浆閲戝瓙宸ュ潑浣跨敤銆?""
    current = st.session_state.get("gold_current_result")
    if isinstance(current, dict):
        return current
    conv = _get_current_conv()
    if conv and isinstance(conv.get("result"), dict):
        return conv["result"]
    return {
        "match": match_data,
        "quality": quality_data,
    }


def _extract_bridge_scores(
    result: dict,
    match_data: Optional[dict] = None,
) -> Optional[dict]:
    """浠庢帰娴嬪櫒缁撴灉鎻愬彇鍒嗘暟锛屼緵宸ュ潑璺宠繃閲嶅 AI 璇勫垎銆?""
    match = match_data if match_data is not None else result.get("match")
    if isinstance(match, dict) and match.get("overall_score") is not None:
        return {
            "overall": float(match.get("overall_score", 0)),
            "star": int(match.get("star_score", 0)),
            "quantify": int(match.get("quant_score", 0)),
            "keyword": int(match.get("keyword_score", 0)),
        }
    quality = result.get("quality")
    if isinstance(quality, dict) and quality.get("overall_score") is not None:
        return {
            "overall": float(quality.get("overall_score", 0)),
            "star": int(quality.get("star_score", 0)),
            "quantify": int(quality.get("quant_score", 0)),
            "keyword": int(quality.get("expression_score", 0)),
        }
    return None


def _bridge_sections_usable(sections: dict) -> bool:
    from engines.resume_parser import sections_look_monolithic

    normalized = {k: str(v or "").strip() for k, v in sections.items()}
    return bool(any(normalized.values())) and not sections_look_monolithic(normalized)


def _navigate_to_workshop(
    result: dict,
    jd_text: str = "",
    match_data: Optional[dict] = None,
) -> None:
    """浠庨噾瀛愭帰娴嬪櫒鎼哄甫鏁版嵁杩涘叆閲戝瓙宸ュ潑銆?""
    conv = _get_current_conv()
    resume = _get_resume_for_followup()
    jd = (jd_text or "").strip()
    if not jd and conv:
        jd = (conv.get("jd") or "").strip()
    if not jd:
        jd = (st.session_state.get("gold_jd_text") or "").strip()

    st.session_state.workshop_resume_text = resume
    st.session_state.workshop_resume_input = resume
    st.session_state.workshop_jd_text = jd
    st.session_state.workshop_jd_input = jd
    st.session_state.workshop_quality_data = result.get("quality")
    st.session_state.workshop_match_data = match_data if match_data is not None else result.get("match")
    st.session_state.workshop_emotion_state = _emotion_state_for_workshop()
    st.session_state.workshop_before_scores = _extract_bridge_scores(
        result, match_data if match_data is not None else result.get("match")
    )

    parsed = result.get("parsed_sections")
    if isinstance(parsed, dict) and _bridge_sections_usable(parsed):
        from engines.resume_parser import SECTION_KEYS

        st.session_state.workshop_sections = {
            key: str(parsed.get(key, "") or "").strip() for key in SECTION_KEYS
        }
        st.session_state.workshop_sections_parsed = True
        st.session_state.workshop_fast_entry = False
    else:
        st.session_state.workshop_sections = {}
        st.session_state.workshop_sections_parsed = False
        st.session_state.workshop_fast_entry = True
    st.session_state.workshop_section_status = {}
    st.session_state.workshop_optimized = {}
    st.session_state.workshop_adopted = {}
    st.session_state.workshop_changes = {}
    st.session_state.workshop_optimize_types = {}
    st.session_state.workshop_manual_editing = None
    st.session_state.workshop_optimize_error = None
    st.session_state.workshop_current_section = "basic_info"

    from ui.sidebar import navigate_to_page

    navigate_to_page("workshop")
    st.rerun()


def _build_followup_prompt(
    resume_text: str,
    result: dict,
    report_text: str,
    question: str,
    history: list[dict],
) -> str:
    analysis_raw = (result.get("analysis") or {}).get("raw_content", "")
    report_raw = (result.get("report") or {}).get("raw_content", "")

    history_lines = []
    for msg in history[-6:]:
        role = "鐢ㄦ埛" if msg.get("role") == "user" else "閲戝瓙"
        history_lines.append(f"{role}锛歿msg.get('content', '')}")

    parts = [
        "銆愬師濮嬬畝鍘嗐€?,
        resume_text or "锛堢畝鍘嗗師鏂囧凡浠庡唴瀛樻竻闄わ紝璇蜂富瑕佷緷鎹笅鏂瑰垎鏋愮粨鏋滃拰鎶ュ憡鍥炵瓟锛?,
        "",
        "銆愬垎鏋愮粨鏋滄憳瑕併€?,
        analysis_raw[:4000] or "鏃?,
        "",
        "銆愭姤鍛婃憳瑕併€?,
        (report_text or report_raw)[:4000] or "鏃?,
    ]
    if history_lines:
        parts.extend(["", "銆愬璇濆巻鍙层€?, *history_lines])
    parts.extend(["", "銆愬綋鍓嶉棶棰樸€?, question])
    return "\n".join(parts)


def _call_followup_analyst(prompt: str) -> str:
    from core.model_router import model_router

    return model_router.call(
        prompt=prompt,
        task_type="resume_followup",
        system_prompt=_FOLLOWUP_SYSTEM_PROMPT,
        max_tokens=800,
    )


def _send_followup_question(question: str, result: dict, report_text: str) -> None:
    question = (question or "").strip()
    if not question:
        return

    history = _get_followup_history()
    resume_text = _get_resume_for_followup()
    prompt = _build_followup_prompt(resume_text, result, report_text, question, history)

    try:
        from ui.error_handler import handle_api_error

        answer = run_with_thinking_chain(
            [
                {"title": "鐞嗚В浣犵殑杩介棶", "desc": "缁撳悎鎶ュ憡涓婁笅鏂囧畾浣嶉棶棰?},
                {"title": "妫€绱㈢畝鍘嗚瘉鎹?, "desc": "浠庡垎鏋愮粨鏋滀腑瀵绘壘鏀拺"},
                {"title": "鐢熸垚閽堝鎬у洖绛?, "desc": "缁欏嚭鍙墽琛岀殑琛ㄨ揪寤鸿"},
            ],
            lambda: _call_followup_analyst(prompt),
            model_name="DeepSeek V3 路 鍒嗘瀽鎺ㄧ悊",
        )
        history.append({"role": "user", "content": question})
        history.append({"role": "assistant", "content": answer.strip()})
    except Exception as e:
        handle_api_error(e, context="gold_followup")


def _render_followup_section(result: dict, report_text: str) -> None:
    """鎶ュ憡涓嬫柟鐨勮拷闂垎鏋愬笀鍖哄煙銆?""
    pending = st.session_state.get("gold_followup_pending")
    if pending:
        st.session_state.gold_followup_pending = None
        _send_followup_question(pending, result, report_text)

    with st.container():
        st.markdown('<div class="gold-followup-wrap">', unsafe_allow_html=True)
        st.markdown('<div class="gold-followup-title">馃挰 杩介棶鍒嗘瀽甯?/div>', unsafe_allow_html=True)

        history = _get_followup_history()
        for msg in history:
            role = msg.get("role", "assistant")
            with st.chat_message("user" if role == "user" else "assistant"):
                st.markdown(msg.get("content", ""))

        q_cols = st.columns(3)
        for i, quick_q in enumerate(_FOLLOWUP_QUICK_QUESTIONS):
            with q_cols[i]:
                if st.button(quick_q, key=f"gold_fq_{i}", use_container_width=True):
                    st.session_state.gold_followup_pending = quick_q
                    st.rerun()

        col_input, col_send = st.columns([5, 1])
        with col_input:
            followup_input = st.text_input(
                "杩介棶杈撳叆",
                placeholder="瀵规姤鍛婃湁鐤戦棶锛熺户缁棶...",
                label_visibility="collapsed",
                key="gold_followup_input",
            )
        with col_send:
            send_clicked = st.button("鍙戦€?, key="gold_followup_send", type="primary", use_container_width=True)

        if send_clicked and followup_input and followup_input.strip():
            _send_followup_question(followup_input.strip(), result, report_text)
            st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)


# ===== 鏍稿績涓氬姟 =====

def _run_detection(resume: str, jd_list: list[dict]) -> dict:
    from ui.error_handler import handle_api_error

    try:
        analysis = run_with_thinking_chain(
            RESUME_ANALYSIS_STEPS,
            lambda: resume_analyzer.analyze(resume),
            model_name="DeepSeek V3 路 鍒嗘瀽鎺ㄧ悊",
        )

        match_results: list[dict] = []
        for jd_item in jd_list:
            content = (jd_item.get("content") or "").strip()
            if not content:
                continue
            name = jd_item.get("name") or "鐩爣宀椾綅"
            analysis_raw = analysis.raw_content

            def _match_jd(jd_content: str = content, raw: str = analysis_raw):
                return JDMatcherV2().match(raw, jd_content)

            match_result = run_with_thinking_chain(
                JD_MATCH_STEPS,
                _match_jd,
                model_name="DeepSeek V3 路 鍒嗘瀽鎺ㄧ悊",
            )
            match_results.append({
                "name": name,
                "result": match_result.model_dump(),
            })

        quality_result = None
        if not match_results:
            quality_result = run_with_thinking_chain(
                RESUME_ANALYSIS_STEPS,
                lambda: ResumeQualityScorer().evaluate(analysis.raw_content),
                model_name="DeepSeek V3 路 鍒嗘瀽鎺ㄧ悊",
            )

        first_raw = ""
        if match_results:
            first_raw = match_results[0]["result"].get("raw_content", "")

        from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeout
        from engines.resume_parser import ResumeParser, heuristic_split_resume, sections_look_monolithic

        def _parse_for_workshop() -> dict[str, str]:
            try:
                sections = ResumeParser().parse(resume).sections
            except Exception as exc:
                logger.warning("[gold_detector] resume parse for workshop failed: %s", exc)
                sections = ResumeParser().parse_fast(resume).sections
            if sections_look_monolithic(sections):
                sections = heuristic_split_resume(resume)
            return sections

        parse_pool = ThreadPoolExecutor(max_workers=1)
        parse_future = parse_pool.submit(_parse_for_workshop)

        report = run_with_thinking_chain(
            [
                {"title": "鏁村悎鍒嗘瀽缁撴灉", "desc": "姹囨€讳紭鍔裤€佸樊璺濅笌缈绘瑙掑害"},
                {"title": "鎾板啓缈绘鎶ュ憡", "desc": "鐢ㄦ俯鏆栬姘斿憟鐜颁綘鐨勬牳蹇冪珵浜夊姏"},
                {"title": "鐢熸垚琛屽姩寤鸿", "desc": "缁欏嚭鍙墽琛岀殑涓嬩竴姝?},
            ],
            lambda: report_generator.generate(
                analysis.raw_content,
                first_raw or "鏃犲矖浣嶄俊鎭紝璇蜂粎鍩轰簬绠€鍘嗗垎鏋?,
            ),
            model_name="DeepSeek V3 路 鍒嗘瀽鎺ㄧ悊",
        )

        parsed_sections: dict[str, str] = {}
        try:
            parsed_sections = parse_future.result(timeout=90)
        except FuturesTimeout:
            logger.warning("[gold_detector] resume parse timeout, using fast fallback")
            parsed_sections = ResumeParser().parse_fast(resume).sections
        except Exception as exc:
            logger.warning("[gold_detector] resume parse error: %s", exc)
            parsed_sections = ResumeParser().parse_fast(resume).sections
        finally:
            parse_pool.shutdown(wait=False, cancel_futures=True)

        recommend_data = None
        try:
            match_data_for_rec = None
            if match_results:
                first_match = match_results[0]["result"]
                if isinstance(first_match, dict):
                    match_data_for_rec = {
                        "keyword_matched": first_match.get("keyword_matched", []),
                        "keyword_missing": first_match.get("keyword_missing", []),
                    }
            recommend_result = JobRecommender().recommend(resume, match_data=match_data_for_rec)
            if recommend_result.recommendations:
                recommend_data = {
                    "recommendations": [
                        {
                            "title": r.title,
                            "match_reason": r.match_reason,
                            "ability_match": r.ability_match,
                            "ability_gap": r.ability_gap,
                            "salary_range": r.salary_range,
                            "search_keyword": r.search_keyword,
                        }
                        for r in recommend_result.recommendations
                    ],
                    "summary": recommend_result.summary,
                }
        except Exception as e:
            logger.warning("[gold_detector] job recommend failed: %s", e)
            recommend_data = {"recommendations": [], "summary": "", "error": True}

        return {
            "analysis": analysis.model_dump(),
            "match_results": match_results,
            "match": match_results[0]["result"] if match_results else None,
            "quality": quality_result.model_dump() if quality_result else None,
            "recommend": recommend_data,
            "report": report.model_dump(),
            "parsed_sections": parsed_sections,
        }
    except Exception as e:
        handle_api_error(e, context="gold_detector")
        raise


def _resolve_jd_content(jd_name: str) -> str:
    """鎸夊矖浣嶅悕绉版煡鎵?JD 鍘熸枃銆?""
    conv = _get_current_conv()
    if conv and conv.get("jd_list"):
        for item in conv["jd_list"]:
            if item.get("name") == jd_name:
                return (item.get("content") or "").strip()
    for item in st.session_state.get("gold_jd_list") or []:
        if item.get("name") == jd_name:
            return (item.get("content") or "").strip()
    return ""


def _save_conversation(
    result: dict,
    resume: str,
    jd: str = "",
    jd_list: Optional[list[dict]] = None,
) -> None:
    new_index = len(st.session_state.gold_conversations) + 1
    conv_id = f"gold_{int(datetime.now().timestamp() * 1000)}"
    conv = {
        "id": conv_id,
        "name": f"鎺㈡祴 {new_index}",
        "resume_snippet": (resume.strip()[:20] + "...") if len(resume.strip()) > 20 else resume.strip(),
        "result": result,
        "resume": resume,
        "jd": jd,
        "jd_list": jd_list or [],
        "created_at": datetime.now().strftime("%m-%d %H:%M"),
    }
    st.session_state.gold_conversations.append(conv)
    # 鎸佷箙鍖栧埌鏂囦欢
    _save_probes(st.session_state.gold_conversations)
    st.session_state.gold_current_conv_id = conv_id
    st.session_state.gold_current_result = result
    st.session_state.gold_show_input = False
    st.session_state.pop("resume_text", None)
    st.session_state.gold_resume_text = ""
    if "gold_resume_input" in st.session_state:
        st.session_state.gold_resume_input = ""


# ===== 鏂版帰娴?=====

def _handle_new_probe() -> None:
    """鐐瑰嚮鏂版帰娴嬶細褰撳墠缁撴灉宸茶嚜鍔ㄤ繚瀛橈紝鐩存帴閲嶇疆杈撳叆鍖?""
    st.session_state.gold_show_input = True
    st.session_state.gold_current_conv_id = None
    st.session_state.gold_current_result = None
    st.session_state.gold_pending_probe = None
    st.session_state.gold_probe_running = False
    st.session_state.gold_resume_text = ""
    st.session_state.gold_jd_text = ""
    st.session_state.gold_jd_list = [{"name": "宀椾綅1", "content": ""}]
    st.session_state.gold_upload_name = None  # 閲嶇疆涓婁紶鏂囦欢鏍囪
    st.session_state.gold_flash_message = None
    # 娓呯┖ uploader widget 鐘舵€侊紝閬垮厤鏃ф枃浠剁姸鎬佽Е鍙戝弽澶嶉噸璺?    if "resume_upload" in st.session_state:
        st.session_state.resume_upload = None
    if "gold_resume_input" in st.session_state:
        st.session_state.gold_resume_input = ""
    if "gold_jd_input" in st.session_state:
        st.session_state.gold_jd_input = ""


# ===== 娓叉煋锛氬乏鍒楀巻鍙查潰鏉?=====

def _render_history_sidebar() -> None:
    """娓叉煋宸﹀垪鍘嗗彶闈㈡澘锛屽缁堟樉绀?""
    gold_convs = st.session_state.get("gold_conversations", [])
    current_conv_id = st.session_state.get("gold_current_conv_id")

    if st.button("锛?鏂版帰娴?, key="gold_history_new", use_container_width=True):
        _handle_new_probe()

    st.markdown('<div class="gold-history-divider"></div>', unsafe_allow_html=True)

    if gold_convs:
        for conv in reversed(gold_convs):
            is_active = conv["id"] == current_conv_id
            name = conv.get("name", "鏈懡鍚?)
            time_str = conv.get("created_at", "")

            # 褰撳墠閫変腑鐨勫姞鍓嶇紑鏍囪
            label = f"鈻?{name} 路 {time_str}" if is_active else f"   {name} 路 {time_str}"

            if st.button(
                label,
                key=f"hist_{conv['id']}",
                use_container_width=True,
            ):
                st.session_state.gold_current_conv_id = conv["id"]
                st.session_state.gold_current_result = conv["result"]
                st.session_state.gold_show_input = False
                st.rerun()
    else:
        st.markdown(
            '<div style="color:#B8AFA5; font-size:12px; padding:8px 4px; line-height:1.6;">杩樻病鏈夋帰娴嬭褰?br/>鎻愪氦绠€鍘嗗悗浼氬嚭鐜板湪杩欓噷</div>',
            unsafe_allow_html=True,
        )


# ===== 娓叉煋锛氬矖浣嶆帹鑽?=====

def _render_job_recommendations_section(result: dict) -> None:
    """娓叉煋宀椾綅鏂瑰悜鎺ㄨ崘锛堟棫璁板綍鏃?recommend 瀛楁鏃惰烦杩囷級銆?""
    recommend_data = result.get("recommend")
    if not recommend_data:
        return

    if recommend_data.get("error") and not recommend_data.get("recommendations"):
        render_section_divider("宀椾綅鎺ㄨ崘")
        st.info("宀椾綅鎺ㄨ崘鏆傛椂涓嶅彲鐢紝璇风◢鍚庨噸璇曘€?)
        return

    if not recommend_data.get("recommendations"):
        return

    recs = [
        JobRecommendation(
            title=r.get("title", ""),
            match_reason=r.get("match_reason", ""),
            ability_match=r.get("ability_match", []),
            ability_gap=r.get("ability_gap", []),
            salary_range=r.get("salary_range", ""),
            search_keyword=r.get("search_keyword", ""),
        )
        for r in recommend_data["recommendations"]
    ]
    rec_result = JobRecommendResult(
        recommendations=recs,
        summary=recommend_data.get("summary", ""),
    )
    render_section_divider("宀椾綅鎺ㄨ崘")
    render_job_recommendations(rec_result)


# ===== 娓叉煋锛氭姤鍛婂尯鍩?=====

def _render_result_block(result: dict) -> None:
    report_text = _parse_report_text(result)
    # 鍏滃簳锛氬鏋滆В鏋愮粨鏋滆繕鏄疛SON寮€澶达紝鍐嶅墺涓€灞?    if report_text and report_text.strip().startswith("{"):
        try:
            parsed = json.loads(report_text)
            if "natural_language_report" in parsed:
                report_text = parsed["natural_language_report"]
        except json.JSONDecodeError:
            pass
    if report_text and report_text.strip().startswith("```"):
        parts = report_text.split("```")
        if len(parts) >= 3:
            inner = parts[1]
            if inner.startswith("json"):
                inner = inner[4:]
            report_text = inner.strip()

    match_score = _parse_match_score(result)
    conv_id = st.session_state.get("gold_current_conv_id")
    conv = _get_current_conv()
    conv_jd = (conv.get("jd") or "") if conv else ""

    match_data = _parse_match_data(result)
    quality_data = _parse_quality_data(result)
    match_results = _get_match_results(result)
    if not quality_data and not match_data and not match_results:
        quality_data = _ensure_quality_data(result, conv_id, conv_jd)

    strengths = _extract_strengths(result, report_text)

    _shell_adapter = EmotionAdapter.from_session()
    st.markdown(f'<div class="{_shell_adapter.get_shell_class()}">', unsafe_allow_html=True)
    render_report_header()
    render_report_body(report_text)

    report_key = str(conv_id or "current")
    if match_results and len(match_results) > 1:
        render_section_divider("鎶曢€掓柟鍚?)
        selected_jd = render_radar_compare(match_results)
        if selected_jd and st.button(
            "馃敤 杩涘叆閲戝瓙宸ュ潑浼樺寲",
            type="primary",
            key=f"radar_workshop_{report_key}",
            use_container_width=True,
        ):
            selected_item = next(m for m in match_results if m["name"] == selected_jd)
            selected_result = selected_item.get("result") or {}
            if hasattr(selected_result, "model_dump"):
                selected_result = selected_result.model_dump()
            jd_content = _resolve_jd_content(selected_jd)
            _navigate_to_workshop(result, jd_text=jd_content, match_data=selected_result)
    elif quality_data and not match_data:
        render_section_divider("璐ㄩ噺璇勪及")
        _render_quality_report(quality_data, report_key=report_key, result=result)
    elif match_data:
        render_section_divider("宀椾綅鍖归厤")
        _render_match_report(match_data, report_key=report_key, result=result)
    elif match_score is not None:
        render_section_divider("鍖归厤搴?)
        st.markdown(f'<div class="gold-score">{match_score}</div>', unsafe_allow_html=True)

    if strengths:
        render_section_divider("浜偣鎻愮偧")
        render_strengths(strengths)

    reframe_pairs = extract_reframe_pairs(result)
    if reframe_pairs:
        render_section_divider("缈绘瀵规瘮")
        render_reframe_compare(reframe_pairs)

    with st.expander("鏌ョ湅璇︾粏鍒嗘瀽鏁版嵁"):
        st.code(json.dumps(result, ensure_ascii=False, indent=2), language="json")

    st.markdown('<div class="gold-export-block">', unsafe_allow_html=True)
    st.markdown('<div class="gold-export-label">瀵煎嚭鎶ュ憡</div>', unsafe_allow_html=True)
    col_exp1, col_exp2, col_exp3 = st.columns(3)

    with col_exp1:
        md_content = _generate_md_report(result, report_text, match_score, strengths, match_data)
        st.download_button(
            label="馃搫 瀵煎嚭 Markdown",
            data=md_content,
            file_name=f"閲戝瓙鎺㈡祴鍣ㄦ姤鍛奯{datetime.now().strftime('%Y%m%d_%H%M')}.md",
            mime="text/markdown",
            key="download_md",
            use_container_width=True,
        )

    with col_exp2:
        docx_bytes = _generate_docx_report(result, report_text, match_score, strengths, match_data)
        st.download_button(
            label="馃摑 瀵煎嚭 Word",
            data=docx_bytes,
            file_name=f"閲戝瓙鎺㈡祴鍣ㄦ姤鍛奯{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx",
            use_container_width=True,
        )

    with col_exp3:
        pdf_bytes = _generate_pdf_report(result, report_text, match_score, strengths, match_data)
        st.download_button(
            label="馃搵 瀵煎嚭 PDF",
            data=pdf_bytes,
            file_name=f"閲戝瓙鎺㈡祴鍣ㄦ姤鍛奯{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
            mime="application/pdf",
            key="download_pdf",
            use_container_width=True,
        )

    st.session_state.gold_report = report_text
    try:
        pdf_bytes = export_gold_report_pdf(report_text)
        st.download_button(
            label="馃摜 涓嬭浇瀹屾暣 PDF 鎶ュ憡",
            data=pdf_bytes,
            file_name="鑱屽満闀滃瓙-閲戝瓙鎺㈡祴鍣ㄦ姤鍛?pdf",
            mime="application/pdf",
            key="download_pdf_core",
            use_container_width=True,
        )
    except Exception as e:
        logger.warning("[gold_detector] PDF export failed: %s", e)

    _render_job_recommendations_section(result)

    st.markdown(
        '<p class="gold-privacy-note">馃敀 绠€鍘嗗師鏂囧凡浠庡唴瀛樹腑娓呴櫎锛屼粎淇濈暀鍒嗘瀽鎶ュ憡</p>',
        unsafe_allow_html=True,
    )

    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    _render_followup_section(result, report_text)

    has_jd = bool(match_data or match_results)
    nav_match = match_data or (match_results[0]["result"] if match_results else None)
    resume_for_bridge = (st.session_state.get("gold_resume_text") or "").strip()
    render_smart_nav(
        get_gold_detector_nav_recommendations(nav_match, has_jd),
        context={
            "strengths": strengths,
            "result": result,
            "resume_snippet": resume_for_bridge,
        },
    )


def _generate_md_report(
    result: dict,
    report_text: str,
    match_score,
    strengths: list,
    match_data: Optional[dict] = None,
) -> str:
    """鐢熸垚Markdown鏍煎紡鎶ュ憡"""
    lines = ["# 閲戝瓙鎺㈡祴鍣?- 绠€鍘嗗垎鏋愭姤鍛奬n"]
    lines.append(f"鐢熸垚鏃堕棿锛歿datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n---\n\n")
    lines.append(f"## 鎶ュ憡姝ｆ枃\n\n{report_text}\n\n")
    if match_data:
        lines.append("## 宀椾綅鍖归厤鎶ュ憡\n\n")
        lines.append(f"- **缁煎悎鍖归厤**锛歿match_data.get('overall_score', match_score)} 鍒哱n")
        lines.append(f"- **鍏抽敭璇嶅尮閰?*锛歿match_data.get('keyword_score', '-')} 鍒哱n")
        lines.append(f"- **STAR 缁撴瀯**锛歿match_data.get('star_score', '-')} 鍒哱n")
        lines.append(f"- **閲忓寲琛ㄨ揪**锛歿match_data.get('quant_score', '-')} 鍒哱n\n")
        if match_data.get("smart_suggestion"):
            lines.append(f"**鏅鸿兘寤鸿**锛歿match_data['smart_suggestion']}\n\n")
    elif match_score is not None:
        lines.append(f"## 鍖归厤搴n\n **{match_score}鍒?*\n\n")
    if strengths:
        lines.append("## 鏍稿績浼樺娍\n\n")
        for s in strengths:
            lines.append(f"- {s}\n")
        lines.append("\n")
    lines.append("---\n*鑱屽満闀滃瓙 路 闄綘璧拌繃鏈€闅剧啲鐨勬眰鑱岃矾*\n")
    return "".join(lines)


def _generate_docx_report(
    result: dict,
    report_text: str,
    match_score,
    strengths: list,
    match_data: Optional[dict] = None,
) -> bytes:
    """鐢熸垚Word鏍煎紡鎶ュ憡"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from io import BytesIO

    doc = Document()

    title = doc.add_heading("閲戝瓙鎺㈡祴鍣?- 绠€鍘嗗垎鏋愭姤鍛?, level=1)
    title.runs[0].font.color.rgb = RGBColor(0x2C, 0x24, 0x20)

    doc.add_paragraph(f"鐢熸垚鏃堕棿锛歿datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    doc.add_heading("鎶ュ憡姝ｆ枃", level=2)
    for para in report_text.split("\n\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(6)

    if match_data:
        doc.add_heading("宀椾綅鍖归厤鎶ュ憡", level=2)
        doc.add_paragraph(f"缁煎悎鍖归厤锛歿match_data.get('overall_score', match_score)} 鍒?)
        doc.add_paragraph(f"鍏抽敭璇嶅尮閰嶏細{match_data.get('keyword_score', '-')} 鍒?)
        doc.add_paragraph(f"STAR 缁撴瀯锛歿match_data.get('star_score', '-')} 鍒?)
        doc.add_paragraph(f"閲忓寲琛ㄨ揪锛歿match_data.get('quant_score', '-')} 鍒?)
        if match_data.get("smart_suggestion"):
            doc.add_paragraph(f"鏅鸿兘寤鸿锛歿match_data['smart_suggestion']}")
    elif match_score is not None:
        doc.add_heading("鍖归厤搴?, level=2)
        score_para = doc.add_paragraph(f"{match_score}鍒?)
        score_para.runs[0].font.size = Pt(20)
        score_para.runs[0].font.color.rgb = RGBColor(0xB8, 0x90, 0x8A)

    if strengths:
        doc.add_heading("鏍稿績浼樺娍", level=2)
        for s in strengths:
            doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("鑱屽満闀滃瓙 路 闄綘璧拌繃鏈€闅剧啲鐨勬眰鑱岃矾").runs[0].font.color.rgb = RGBColor(
        0x9E, 0x8E, 0x83
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf_report(
    result: dict,
    report_text: str,
    match_score,
    strengths: list,
    match_data: Optional[dict] = None,
) -> bytes:
    """鐢熸垚PDF鏍煎紡鎶ュ憡锛堜娇鐢╳easyprint HTML杞琍DF锛岃В鍐充腑鏂囦贡鐮侊級"""
    from io import BytesIO

    # 鏋勫缓HTML鍐呭
    html_parts = [
        '<html><head><meta charset="utf-8">',
        '<style>body{font-family: "Microsoft YaHei", "SimHei", sans-serif; padding: 40px; color: #2C2420; line-height: 1.8;}',
        'h1{color: #2C2420; border-bottom: 2px solid #B8908A; padding-bottom: 10px;}',
        'h2{color: #2C2420; margin-top: 24px;}',
        '.score{color: #B8908A; font-size: 28px; font-weight: bold;}',
        '.footer{color: #9E8E83; font-size: 12px; margin-top: 40px; border-top: 1px solid #ddd; padding-top: 10px;}',
        'ul{padding-left: 20px;} li{margin-bottom: 6px;}</style></head><body>',
        '<h1>閲戝瓙鎺㈡祴鍣?- 绠€鍘嗗垎鏋愭姤鍛?/h1>',
        f'<p>鐢熸垚鏃堕棿锛歿datetime.now().strftime("%Y-%m-%d %H:%M")}</p>',
        '<h2>鎶ュ憡姝ｆ枃</h2>',
    ]

    for para in report_text.split("\n\n"):
        if para.strip():
            html_parts.append(f'<p>{para.strip().replace(chr(10), "<br/>")}</p>')

    if match_data:
        html_parts.append("<h2>宀椾綅鍖归厤鎶ュ憡</h2>")
        html_parts.append(f"<p>缁煎悎鍖归厤锛歿match_data.get('overall_score', match_score)} 鍒?/p>")
        html_parts.append(f"<p>鍏抽敭璇嶅尮閰嶏細{match_data.get('keyword_score', '-')} 鍒?/p>")
        html_parts.append(f"<p>STAR 缁撴瀯锛歿match_data.get('star_score', '-')} 鍒?/p>")
        html_parts.append(f"<p>閲忓寲琛ㄨ揪锛歿match_data.get('quant_score', '-')} 鍒?/p>")
        if match_data.get("smart_suggestion"):
            html_parts.append(f"<p>鏅鸿兘寤鸿锛歿match_data['smart_suggestion']}</p>")
    elif match_score is not None:
        html_parts.append(f'<h2>鍖归厤搴?/h2><p class="score">{match_score}鍒?/p>')

    if strengths:
        html_parts.append('<h2>鏍稿績浼樺娍</h2><ul>')
        for s in strengths:
            html_parts.append(f'<li>{s}</li>')
        html_parts.append('</ul>')

    html_parts.append('<p class="footer">鑱屽満闀滃瓙 路 闄綘璧拌繃鏈€闅剧啲鐨勬眰鑱岃矾</p>')
    html_parts.append('</body></html>')

    html_content = "".join(html_parts)

    try:
        from weasyprint import HTML as WeasyHTML
        pdf_bytes = WeasyHTML(string=html_content).write_pdf()
        return pdf_bytes
    except Exception:
        # weasyprint鍦╓indows鍙兘鍥犵郴缁熷簱缂哄け鎶汷SError锛岃繖閲岀粺涓€闄嶇骇
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.colors import HexColor
            from io import BytesIO

            buf = BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4)
            styles = getSampleStyleSheet()

            # 灏濊瘯娉ㄥ唽涓枃瀛椾綋
            font_name = "ChineseFont"
            font_registered = False
            for font_path in [
                "C:/Windows/Fonts/msyh.ttc",   # 寰蒋闆呴粦
                "C:/Windows/Fonts/simhei.ttf",  # 榛戜綋
                "C:/Windows/Fonts/simsun.ttc",  # 瀹嬩綋
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # Linux
                "/System/Library/Fonts/PingFang.ttc",  # macOS
            ]:
                if Path(font_path).exists():
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        font_registered = True
                        break
                    except Exception:
                        continue

            body_font = font_name if font_registered else "Helvetica"
            title_font = font_name if font_registered else "Helvetica"

            title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontName=title_font, textColor=HexColor("#2C2420"), fontSize=18, spaceAfter=12)
            heading_style = ParagraphStyle("CustomHeading", parent=styles["Heading2"], fontName=title_font, textColor=HexColor("#2C2420"), fontSize=14, spaceAfter=8)
            body_style = ParagraphStyle("CustomBody", parent=styles["Normal"], fontName=body_font, textColor=HexColor("#2C2420"), fontSize=10, leading=16, spaceAfter=6)

            elements = []
            elements.append(Paragraph("Gold Detector - Resume Report", title_style))
            elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
            elements.append(Spacer(1, 12))

            elements.append(Paragraph("Report", heading_style))
            for para in report_text.split("\n\n"):
                if para.strip():
                    safe_text = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
                    elements.append(Paragraph(safe_text, body_style))

            if match_score is not None:
                elements.append(Paragraph("Match Score", heading_style))
                score_style = ParagraphStyle("Score", parent=body_style, fontSize=20, textColor=HexColor("#B8908A"))
                elements.append(Paragraph(f"{match_score}", score_style))

            if strengths:
                elements.append(Paragraph("Core Strengths", heading_style))
                for s in strengths:
                    safe_s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    elements.append(Paragraph(f"鈥?{safe_s}", body_style))

            elements.append(Spacer(1, 20))
            footer_style = ParagraphStyle("Footer", parent=body_style, textColor=HexColor("#9E8E83"), fontSize=8)
            elements.append(Paragraph("Workplace Mirror", footer_style))

            doc.build(elements)
            return buf.getvalue()
        except Exception as e:
            # 鏈€缁堥檷绾э細杩斿洖HTML鏂囦欢浣嗘敼鍚嶄负PDF
            return html_content.encode("utf-8")


# ===== 娓叉煋锛氳緭鍏ュ尯 =====

def _render_input_area() -> None:
    adapter = EmotionAdapter.from_session()
    if adapter.get_layout_mode() == "single_column":
        st.info("馃挕 涓€娆＄湅涓€涓澘鍧楀氨濂斤紝涓嶆€ャ€傛帰娴嬪畬鎴愬悗鎶ュ憡浼氫负浣犵畝鍖栧睍绀恒€?)

    uploaded_file = st.file_uploader(
        "涓婁紶绠€鍘嗭紙鏀寔PDF锛?,
        type=["pdf"],
        key="resume_upload",
    )
    if uploaded_file and st.session_state.gold_upload_name != uploaded_file.name:
        with st.spinner("姝ｅ湪瑙ｆ瀽PDF..."):
            try:
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pdfminer")
                logging.getLogger("pdfminer.pdffont").setLevel(logging.ERROR)`nimport pdfplumber

                pdf_bytes = io.BytesIO(uploaded_file.read())
                with pdfplumber.open(pdf_bytes) as pdf:
                    resume_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                st.session_state.gold_resume_text = resume_text
                st.session_state.gold_upload_name = uploaded_file.name
                if "gold_resume_input" in st.session_state:
                    st.session_state.gold_resume_input = resume_text
            except ModuleNotFoundError:
                st.error("PDF瑙ｆ瀽渚濊禆鏈畨瑁咃細璇锋墽琛?`python -m pip install pdfplumber` 鍚庨噸璇曘€?)
            except Exception as exc:
                st.error(f"PDF瑙ｆ瀽澶辫触锛歿exc}")

    resume = st.text_area(
        "绠€鍘嗗唴瀹?,
        placeholder="鎶婁綘鐨勭畝鍘嗙矘璐村埌杩欓噷锛屾垨涓婁紶PDF鏂囦欢...",
        height=250,
        key="gold_resume_input",
    )

    st.markdown("### 馃搵 宀椾綅鎻忚堪锛堝彲娣诲姞澶氫釜瀵规瘮锛?)

    for i, jd_item in enumerate(st.session_state.gold_jd_list):
        col_name, col_content, col_remove = st.columns([1, 6, 1])
        with col_name:
            jd_item["name"] = st.text_input(
                "宀椾綅鍚嶇О",
                value=jd_item.get("name", f"宀椾綅{i + 1}"),
                key=f"jd_name_{i}",
                label_visibility="collapsed",
                placeholder=f"宀椾綅{i + 1}鍚嶇О",
            )
        with col_content:
            jd_item["content"] = st.text_area(
                "JD鍐呭",
                value=jd_item.get("content", ""),
                height=120,
                key=f"jd_content_{i}",
                label_visibility="collapsed",
                placeholder=f"绮樿创宀椾綅{i + 1}鐨凧D...",
            )
        with col_remove:
            if len(st.session_state.gold_jd_list) > 1:
                if st.button("鉁?, key=f"jd_remove_{i}"):
                    st.session_state.gold_jd_list.pop(i)
                    st.rerun()

    if len(st.session_state.gold_jd_list) < 3:
        if st.button("锛?娣诲姞宀椾綅瀵规瘮", key="add_jd"):
            idx = len(st.session_state.gold_jd_list) + 1
            st.session_state.gold_jd_list.append({"name": f"宀椾綅{idx}", "content": ""})
            st.rerun()

    st.session_state.gold_resume_text = st.session_state.get("gold_resume_input", "")
    jd_text = st.session_state.gold_jd_list[0]["content"] if st.session_state.gold_jd_list else ""
    st.session_state.gold_jd_text = jd_text

    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("寮€濮嬫帰娴?, type="primary"):
            if not resume.strip():
                st.warning("璇峰厛杈撳叆鎴栦笂浼犵畝鍘?)
            else:
                st.session_state.gold_pending_probe = {
                    "resume": resume,
                    "jd_list": list(st.session_state.gold_jd_list),
                }
                st.session_state.gold_probe_running = True
                st.rerun()
    with col2:
        if st.button("娓呯┖"):
            st.session_state.gold_resume_text = ""
            st.session_state.gold_jd_text = ""
            st.session_state.gold_jd_list = [{"name": "宀椾綅1", "content": ""}]
            st.session_state.gold_upload_name = None
            if "gold_resume_input" in st.session_state:
                st.session_state.gold_resume_input = ""
            st.rerun()


# ===== 娓叉煋锛氬彸鍒椾富鍐呭 =====

def _render_main_content() -> None:
    """娓叉煋鍙冲垪涓诲唴瀹瑰尯"""
    # ===== 濡傛灉鏈夊緟鎵ц鐨勬帰娴嬶紝鍦ㄨ繖閲岃窇锛堣繘搴﹀湪鍙冲垪鍙锛?=====
    if st.session_state.get("gold_pending_probe"):
        probe_data = st.session_state.gold_pending_probe
        st.session_state.gold_probe_running = True

        st.markdown("### 馃攳 姝ｅ湪鍒嗘瀽浣犵殑绠€鍘?..")
        st.markdown(
            '<div style="color:#8C8279; font-size:13px;">棰勮闇€瑕?-2鍒嗛挓锛岃鑰愬績绛夊緟</div>',
            unsafe_allow_html=True,
        )
        st.info("鎺㈡祴浠诲姟宸插惎鍔紝姝ｅ湪鎵ц涓紝璇峰嬁閲嶅鐐瑰嚮鎴栧埛鏂伴〉闈€?)

        try:
            jd_list = probe_data.get("jd_list")
            if jd_list is None:
                legacy_jd = probe_data.get("jd", "")
                jd_list = [{"name": "宀椾綅1", "content": legacy_jd}] if legacy_jd else [{"name": "宀椾綅1", "content": ""}]

            result = _run_detection(probe_data["resume"], jd_list)
            primary_jd = (jd_list[0].get("content") or "") if jd_list else ""
            _save_conversation(result, probe_data["resume"], primary_jd, jd_list=jd_list)
            st.session_state.gold_pending_probe = None
            st.session_state.gold_flash_message = "鎺㈡祴瀹屾垚锛?
            st.session_state.gold_flash_type = "success"
            st.rerun()
            return
        except Exception as e:
            from ui.error_handler import get_friendly_message

            logger.warning("[gold_detector] %s: %s", type(e).__name__, e)
            st.session_state.gold_pending_probe = None
            st.session_state.gold_flash_message = get_friendly_message(e)
            st.session_state.gold_flash_type = "error"
            st.session_state.gold_show_input = True
            st.rerun()
            return
        finally:
            st.session_state.gold_probe_running = False

    # ===== Flash 娑堟伅 =====
    flash_msg = st.session_state.get("gold_flash_message")
    if flash_msg:
        flash_type = st.session_state.get("gold_flash_type", "info")
        if flash_type == "success":
            st.success(flash_msg)
        elif flash_type == "error":
            st.error(flash_msg)
        else:
            st.info(flash_msg)
        st.session_state.gold_flash_message = None

    # ===== 姝ｅ父鍐呭 =====
    if st.session_state.get("gold_show_input"):
        _render_input_area()
        return

    conv_id = st.session_state.get("gold_current_conv_id")
    conv = next((c for c in st.session_state.gold_conversations if c["id"] == conv_id), None)
    if conv:
        st.session_state.gold_current_result = conv["result"]
        _render_result_block(conv["result"])
    else:
        current_result = st.session_state.get("gold_current_result")
        if current_result:
            _render_result_block(current_result)
        else:
            st.session_state.gold_show_input = True
            _render_input_area()


# ===== 涓诲叆鍙?=====

def render():
    track_module_enter("閲戝瓙鎺㈡祴鍣?)
    _inject_styles()
    _init_state()

    render_page_header("閲戝瓙鎺㈡祴鍣?, "鎶婄粡鍘嗘斁涓婃潵锛岀湅鐪嬩綘鐨勬牳蹇冪珵浜夊姏")
    apply_emotion_breath()

    # 涓讳綋锛氬乏鍒楀巻鍙诧紙绐勶級 + 鍙冲垪鍐呭锛堝锛?    col_history, col_content = st.columns([1, 7], gap="small")

    with col_history:
        _render_history_sidebar()

    with col_content:
        _render_main_content()

